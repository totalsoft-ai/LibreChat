#!/usr/bin/env python3
"""
Parse a Word (.docx) document and generate per-section Markdown files + images.

Usage:
    python parse_docx.py <input.docx> <images_dir> <output_dir> <lang>

Output:
    - <output_dir>/sections.json   — metadata for all sections
    - <output_dir>/<id>.md         — one file per top-level heading
    - <images_dir>/<lang>_<n>.png  — extracted images
"""

import sys
import os
import json
import re
import io
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    Image = None
    PIL_AVAILABLE = False

# Sections whose titles match these patterns will be skipped entirely
SKIP_SECTION_PATTERNS = [
    r'^document\s+version',
    r'^version\s+history',
    r'^revision\s+history',
    r'^document\s+information',
    r'^document\s+control',
    r'^table\s+of\s+contents',
    r'^cuprins',
    r'^istoricul\s+document',
    r'^versiunea\s+document',
]

def _should_skip_section(title: str) -> bool:
    t = title.strip().lower()
    for pat in SKIP_SECTION_PATTERNS:
        if re.search(pat, t):
            return True
    return False


def slugify(text):
    """Convert heading text to a URL-safe slug."""
    text = text.lower().strip()
    # Transliterate common Romanian chars
    text = text.replace('ă', 'a').replace('â', 'a').replace('î', 'i')
    text = text.replace('ș', 's').replace('ş', 's')
    text = text.replace('ț', 't').replace('ţ', 't')
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[\s_-]+', '-', text)
    text = re.sub(r'^-+|-+$', '', text)
    return text or 'section'


def extract_image(rel, images_dir, lang, image_counter):
    """Extract an image from a relationship and save it. Returns the image filename or None."""
    try:
        image_data = rel.target_part.blob
        content_type = rel.target_part.content_type

        ext_map = {
            'image/png': 'png',
            'image/jpeg': 'jpg',
            'image/gif': 'gif',
            'image/bmp': 'bmp',
            'image/tiff': 'tiff',
            'image/wmf': 'wmf',
            'image/emf': 'emf',
            'image/x-wmf': 'wmf',
            'image/x-emf': 'emf',
        }
        ext = ext_map.get(content_type, 'png')
        is_vector = ext in ('wmf', 'emf')

        # Final filename always .png
        filename = f'{lang}_{image_counter:04d}.png'
        filepath = os.path.join(images_dir, filename)

        if is_vector:
            if PIL_AVAILABLE:
                try:
                    img = Image.open(io.BytesIO(image_data))
                    img.save(filepath, 'PNG')
                    return filename
                except Exception:
                    pass
            # Cannot convert vector image without PIL — skip it
            return None
        else:
            # For jpeg, keep original extension so browsers render correctly
            if ext == 'jpg':
                filename = f'{lang}_{image_counter:04d}.jpg'
                filepath = os.path.join(images_dir, filename)
            with open(filepath, 'wb') as f:
                f.write(image_data)
            return filename

    except Exception as e:
        print(f"Warning: could not extract image: {e}", file=sys.stderr)
        return None


def get_heading_level(paragraph):
    """Return heading level (1-6) if it's a heading, else None."""
    style_name = paragraph.style.name.lower()
    if style_name.startswith('heading'):
        parts = style_name.split()
        try:
            return int(parts[-1])
        except (ValueError, IndexError):
            pass
    if paragraph.style.name == 'Title':
        return 1
    return None


def runs_to_markdown(runs):
    """
    Convert a list of text-only runs to markdown, merging consecutive same-format runs
    so that bold/italic markers don't appear as literal text due to fragmented runs.
    """
    # Group consecutive runs by (bold, italic)
    groups = []  # list of ((bold, italic), text)
    for run in runs:
        text = run.text
        if not text:
            continue
        fmt = (bool(run.bold), bool(run.italic))
        if groups and groups[-1][0] == fmt:
            # Merge with previous group
            groups[-1] = (fmt, groups[-1][1] + text)
        else:
            groups.append((fmt, text))

    result = []
    for (bold, italic), text in groups:
        if not text:
            continue
        stripped = text.strip()
        if not stripped:
            # Whitespace-only run — preserve it outside any markers
            result.append(text)
            continue
        leading = text[:len(text) - len(text.lstrip())]
        trailing = text[len(text.rstrip()):]
        if bold and italic:
            result.append(f'{leading}***{stripped}***{trailing}')
        elif bold:
            result.append(f'{leading}**{stripped}**{trailing}')
        elif italic:
            result.append(f'{leading}*{stripped}*{trailing}')
        else:
            result.append(text)
    return ''.join(result)


def paragraph_to_markdown(paragraph, images_dir, lang, image_counter_ref, image_rels):
    """Convert a paragraph to Markdown. Returns markdown string (may be empty)."""
    level = get_heading_level(paragraph)

    # Collect inline content: images and text runs (grouped by position)
    # We process runs in order; images are inline within runs
    text_runs = []   # accumulates runs between images
    parts = []       # final parts list (strings)

    for run in paragraph.runs:
        image_found = False
        for elem in run._r:
            if elem.tag.endswith('}drawing') or elem.tag.endswith('}pict'):
                blips = elem.findall('.//' + qn('a:blip'))
                for blip in blips:
                    embed = blip.get(qn('r:embed'))
                    if embed and embed in image_rels:
                        # Flush accumulated text runs first
                        if text_runs:
                            parts.append(runs_to_markdown(text_runs))
                            text_runs = []
                        filename = extract_image(
                            image_rels[embed], images_dir, lang, image_counter_ref[0]
                        )
                        if filename:
                            image_counter_ref[0] += 1
                            parts.append(f'\n\n![image](/help-images/{filename})\n\n')
                        image_found = True
        if not image_found:
            text_runs.append(run)

    # Flush remaining text runs
    if text_runs:
        parts.append(runs_to_markdown(text_runs))

    text = ''.join(parts).strip()

    if not text:
        return ''

    if level:
        prefix = '#' * level
        # For headings, strip any markdown markers that leaked in
        clean_text = text.replace('**', '').replace('*', '').strip()
        return f'{prefix} {clean_text}'

    # List items
    style_name = paragraph.style.name.lower()
    if 'list bullet' in style_name or paragraph.style.name.startswith('List Bullet'):
        return f'- {text}'
    if 'list number' in style_name or paragraph.style.name.startswith('List Number'):
        return f'1. {text}'

    return text


def table_to_markdown(table):
    """Convert a table to Markdown."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        rows.append('| ' + ' | '.join(cells) + ' |')
        if i == 0:
            rows.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
    return '\n'.join(rows)


def parse_document(docx_path, images_dir, output_dir, lang):
    """Parse the document and write per-section markdown files."""
    doc = Document(docx_path)
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(images_dir, exist_ok=True)

    # Build relationship map for images
    image_rels = {}
    for rel in doc.part.rels.values():
        if 'image' in rel.reltype:
            image_rels[rel.rId] = rel

    image_counter = [1]

    # Parse all blocks (paragraphs + tables) in document order
    blocks = []
    body = doc.element.body
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p' and child in para_map:
            blocks.append(('paragraph', para_map[child]))
        elif tag == 'tbl' and child in table_map:
            blocks.append(('table', table_map[child]))

    # Detect split level: use H1 if present, else fall back to H2
    h1_count = sum(
        1 for bt, b in blocks
        if bt == 'paragraph' and get_heading_level(b) == 1
    )
    split_level = 1 if h1_count > 0 else 2
    if split_level == 2:
        print(f"Warning: no H1 headings found in {docx_path}, splitting on H2", file=sys.stderr)

    # Split into sections by top-level headings
    sections = []       # list of (title, lines[])
    current_section = None
    current_lines = []

    for block_type, block in blocks:
        if block_type == 'paragraph':
            level = get_heading_level(block)
            if level == split_level:
                if current_section is not None:
                    sections.append((current_section, current_lines))
                heading_text = block.text.strip()
                current_section = heading_text
                current_lines = [f'# {heading_text}']
            else:
                md = paragraph_to_markdown(block, images_dir, lang, image_counter, image_rels)
                if md:
                    current_lines.append(md)
        elif block_type == 'table':
            md = table_to_markdown(block)
            if md:
                current_lines.append('\n' + md + '\n')

    if current_section is not None:
        sections.append((current_section, current_lines))

    # Write files and build metadata (skip unwanted sections)
    seen_slugs = {}
    sections_meta = []
    order = 0

    for title, lines in sections:
        if _should_skip_section(title):
            print(f"  Skipping section: {title!r}", file=sys.stderr)
            continue

        slug = slugify(title)
        if slug in seen_slugs:
            seen_slugs[slug] += 1
            slug = f'{slug}-{seen_slugs[slug]}'
        else:
            seen_slugs[slug] = 0

        content = '\n\n'.join(line for line in lines if line)
        filepath = os.path.join(output_dir, f'{slug}.md')
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content + '\n')

        sections_meta.append({
            'id': slug,
            'title': title,
            'order': order,
            'filename': f'{slug}.md',
        })
        order += 1

    meta_path = os.path.join(output_dir, 'sections.json')
    with open(meta_path, 'w', encoding='utf-8') as f:
        json.dump(
            {
                'sections': sections_meta,
                'default': sections_meta[0]['id'] if sections_meta else 'index',
            },
            f,
            ensure_ascii=False,
            indent=2,
        )

    print(f"Generated {len(sections_meta)} sections in {output_dir}", file=sys.stderr)
    return sections_meta


if __name__ == '__main__':
    if len(sys.argv) != 5:
        print(f"Usage: {sys.argv[0]} <input.docx> <images_dir> <output_dir> <lang>", file=sys.stderr)
        sys.exit(1)

    docx_path, images_dir, output_dir, lang = sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4]

    if not os.path.isfile(docx_path):
        print(f"Error: file not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    sections = parse_document(docx_path, images_dir, output_dir, lang)
    print(json.dumps({'sections': sections}))
